"""Document extractor for EREN Knowledge Ingestion Pipeline.

Extracts text from various document formats.
"""

from __future__ import annotations

import time
import uuid
from abc import ABC, abstractmethod
from typing import TYPE_CHECKING

from core.ingestion.types import (
    DocumentType,
    RawDocument,
    ExtractedDocument,
    IngestionMetadata,
)

if TYPE_CHECKING:
    pass


class BaseExtractor(ABC):
    """Abstract base class for document extractors."""

    @abstractmethod
    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from document.

        Args:
            raw: Raw document.

        Returns:
            Extracted document.
        """
        pass

    @property
    @abstractmethod
    def supported_types(self) -> list[DocumentType]:
        """Get supported document types."""
        pass


class PDFExtractor(BaseExtractor):
    """PDF document extractor."""

    @property
    def supported_types(self) -> list[DocumentType]:
        """Get supported types."""
        return [DocumentType.PDF]

    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from PDF.

        Args:
            raw: Raw PDF document.

        Returns:
            Extracted document.
        """
        start_time = time.time()
        warnings = []

        try:
            # Try to use PyPDF2 or pdfplumber
            text = self._extract_pdf_text(raw.content)
        except ImportError:
            # Fallback: try simple text extraction
            try:
                text = self._extract_simple(raw.content)
            except Exception as e:
                text = ""
                warnings.append(f"Extraction warning: {str(e)}")

        extraction_time = int((time.time() - start_time) * 1000)

        return ExtractedDocument(
            document_id=str(uuid.uuid4()),
            content=text,
            document_type=DocumentType.PDF,
            metadata=raw.metadata,
            extraction_time_ms=extraction_time,
            warnings=warnings,
        )

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using PyPDF2."""
        try:
            import PyPDF2
            from io import BytesIO

            pdf_file = BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return "\n\n".join(text_parts)

        except ImportError:
            raise ImportError("PyPDF2 not installed")

    def _extract_simple(self, content: bytes) -> str:
        """Simple text extraction fallback."""
        # Try to decode as text
        try:
            return content.decode("utf-8", errors="ignore")
        except Exception:
            return ""


class DocxExtractor(BaseExtractor):
    """DOCX document extractor."""

    @property
    def supported_types(self) -> list[DocumentType]:
        """Get supported types."""
        return [DocumentType.DOCX]

    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from DOCX.

        Args:
            raw: Raw DOCX document.

        Returns:
            Extracted document.
        """
        start_time = time.time()
        warnings = []

        try:
            import docx
            from io import BytesIO

            doc_file = BytesIO(raw.content)
            doc = docx.Document(doc_file)

            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)

        except ImportError:
            warnings.append("python-docx not installed, using fallback")
            try:
                text = content.decode("utf-8", errors="ignore")
            except Exception:
                text = ""

        extraction_time = int((time.time() - start_time) * 1000)

        return ExtractedDocument(
            document_id=str(uuid.uuid4()),
            content=text,
            document_type=DocumentType.DOCX,
            metadata=raw.metadata,
            extraction_time_ms=extraction_time,
            warnings=warnings,
        )


class TextExtractor(BaseExtractor):
    """Plain text extractor."""

    @property
    def supported_types(self) -> list[DocumentType]:
        """Get supported types."""
        return [DocumentType.TXT, DocumentType.MARKDOWN]

    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from plain text.

        Args:
            raw: Raw text document.

        Returns:
            Extracted document.
        """
        start_time = time.time()

        # Try to decode
        try:
            text = raw.content.decode("utf-8", errors="replace")
        except Exception:
            text = str(raw.content)

        extraction_time = int((time.time() - start_time) * 1000)

        return ExtractedDocument(
            document_id=str(uuid.uuid4()),
            content=text,
            document_type=raw.document_type,
            metadata=raw.metadata,
            extraction_time_ms=extraction_time,
        )


class HTMLExtractor(BaseExtractor):
    """HTML document extractor."""

    @property
    def supported_types(self) -> list[DocumentType]:
        """Get supported types."""
        return [DocumentType.HTML]

    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from HTML.

        Args:
            raw: Raw HTML document.

        Returns:
            Extracted document.
        """
        start_time = time.time()
        warnings = []

        try:
            from bs4 import BeautifulSoup

            html_text = raw.content.decode("utf-8", errors="ignore")
            soup = BeautifulSoup(html_text, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text(separator="\n", strip=True)

        except ImportError:
            warnings.append("BeautifulSoup not installed, using fallback")
            text = raw.content.decode("utf-8", errors="ignore")

        extraction_time = int((time.time() - start_time) * 1000)

        return ExtractedDocument(
            document_id=str(uuid.uuid4()),
            content=text,
            document_type=DocumentType.HTML,
            metadata=raw.metadata,
            extraction_time_ms=extraction_time,
            warnings=warnings,
        )


class FHIRExtractor(BaseExtractor):
    """FHIR resource extractor."""

    @property
    def supported_types(self) -> list[DocumentType]:
        """Get supported types."""
        return [DocumentType.FHIR]

    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from FHIR resources.

        Args:
            raw: Raw FHIR document.

        Returns:
            Extracted document.
        """
        start_time = time.time()
        warnings = []

        try:
            import json

            # Parse FHIR JSON
            fhir_data = json.loads(raw.content.decode("utf-8"))

            # Extract relevant text
            text_parts = []

            # Handle different FHIR resource types
            if isinstance(fhir_data, dict):
                resource_type = fhir_data.get("resourceType", "Unknown")

                # Extract text from various fields
                if "text" in fhir_data and "div" in fhir_data["text"]:
                    text_parts.append(fhir_data["text"]["div"])

                if "contentString" in fhir_data:
                    text_parts.append(fhir_data["contentString"])

                # Add specific fields based on resource type
                if resource_type == "Patient":
                    if "name" in fhir_data:
                        for name in fhir_data["name"]:
                            text_parts.append(f"Patient: {name.get('text', '')}")

                elif resource_type == "Observation":
                    if "code" in fhir_data:
                        text_parts.append(f"Observation: {fhir_data['code'].get('text', '')}")

                elif resource_type == "Condition":
                    if "code" in fhir_data:
                        text_parts.append(f"Condition: {fhir_data['code'].get('text', '')}")

            text = "\n\n".join(text_parts)

        except json.JSONDecodeError:
            warnings.append("Invalid FHIR JSON, treating as plain text")
            text = raw.content.decode("utf-8", errors="ignore")

        extraction_time = int((time.time() - start_time) * 1000)

        return ExtractedDocument(
            document_id=str(uuid.uuid4()),
            content=text,
            document_type=DocumentType.FHIR,
            metadata=raw.metadata,
            extraction_time_ms=extraction_time,
            warnings=warnings,
        )


class HL7Extractor(BaseExtractor):
    """HL7 message extractor."""

    @property
    def supported_types(self) -> list[DocumentType]:
        """Get supported types."""
        return [DocumentType.HL7]

    async def extract(self, raw: RawDocument) -> ExtractedDocument:
        """Extract text from HL7 messages.

        Args:
            raw: Raw HL7 document.

        Returns:
            Extracted document.
        """
        start_time = time.time()
        warnings = []

        try:
            content = raw.content.decode("utf-8", errors="ignore")
            lines = content.split("\r\n")

            text_parts = []

            for line in lines:
                if not line.strip():
                    continue

                # Parse HL7 segments
                segments = line.split("|")
                if segments:
                    segment_type = segments[0]

                    # Extract relevant information based on segment type
                    if segment_type == "PID":
                        # Patient identification
                        if len(segments) > 5:
                            text_parts.append(f"Patient: {segments[5]}")

                    elif segment_type == "OBR":
                        # Observation request
                        if len(segments) > 4:
                            text_parts.append(f"Observation: {segments[4]}")

                    elif segment_type == "OBX":
                        # Observation result
                        if len(segments) > 5:
                            text_parts.append(f"Result: {segments[5]}")

                    elif segment_type == "TXA":
                        # Transcription
                        if len(segments) > 7:
                            text_parts.append(f"Transcription: {segments[7]}")

            text = "\n\n".join(text_parts)

        except Exception as e:
            warnings.append(f"HL7 parsing warning: {str(e)}")
            text = raw.content.decode("utf-8", errors="ignore")

        extraction_time = int((time.time() - start_time) * 1000)

        return ExtractedDocument(
            document_id=str(uuid.uuid4()),
            content=text,
            document_type=DocumentType.HL7,
            metadata=raw.metadata,
            extraction_time_ms=extraction_time,
            warnings=warnings,
        )


# Factory for extractors
class ExtractorFactory:
    """Factory for document extractors."""

    _extractors: dict[DocumentType, BaseExtractor] = {}

    @classmethod
    def register(cls, extractor: BaseExtractor) -> None:
        """Register an extractor.

        Args:
            extractor: Extractor to register.
        """
        for doc_type in extractor.supported_types:
            cls._extractors[doc_type] = extractor

    @classmethod
    def get(cls, document_type: DocumentType) -> BaseExtractor | None:
        """Get extractor for document type.

        Args:
            document_type: Document type.

        Returns:
            Extractor or None.
        """
        return cls._extractors.get(document_type)

    @classmethod
    def get_for_document(cls, raw: RawDocument) -> BaseExtractor | None:
        """Get extractor for raw document.

        Args:
            raw: Raw document.

        Returns:
            Extractor or None.
        """
        # Try by document type
        extractor = cls._extractors.get(raw.document_type)
        if extractor:
            return extractor

        # Try by filename extension
        if raw.filename:
            ext = raw.filename.lower().split(".")[-1]
            type_map = {
                "pdf": DocumentType.PDF,
                "docx": DocumentType.DOCX,
                "doc": DocumentType.DOCX,
                "txt": DocumentType.TXT,
                "md": DocumentType.MARKDOWN,
                "html": DocumentType.HTML,
                "htm": DocumentType.HTML,
            }
            doc_type = type_map.get(ext, DocumentType.UNKNOWN)
            return cls._extractors.get(doc_type)

        return None

    @classmethod
    def list_supported_types(cls) -> list[DocumentType]:
        """List all supported document types.

        Returns:
            List of supported types.
        """
        return list(cls._extractors.keys())

    @classmethod
    def register_defaults(cls) -> None:
        """Register default extractors."""
        cls.register(PDFExtractor())
        cls.register(DocxExtractor())
        cls.register(TextExtractor())
        cls.register(HTMLExtractor())
        cls.register(FHIRExtractor())
        cls.register(HL7Extractor())
